"""Tests for the DOCX + Markdown report writer.

The writer is responsible for the canonical revision-plan deliverable.
These tests ensure the new layout (executive summary, per-KP cheat-sheets,
sensitivity, appendices) is correctly generated and that the bilingual
behaviour matches the spec.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from scripts.pattern_coverage import PatternCoverage, PatternOccurrence
from scripts.report_writer import (
    APPENDIX_AUDIT,
    APPENDIX_METHODOLOGY,
    APPENDIX_PATTERNS,
    SECTION_CHEAT_SHEETS,
    SECTION_EXEC_SUMMARY,
    SECTION_PREDICTIONS,
    SECTION_SENSITIVITY,
    write_docx,
    write_markdown,
)
from scripts.sensitivity import (
    SensitivityCell,
    SensitivitySweep,
)
from scripts.statistical_model import KPPosterior


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def _posterior(
    kp_id: str = "L13.03",
    tier: str = "anchor",
    posterior_mean: float = 0.85,
    ci_lower_95: float = 0.55,
    ci_upper_95: float = 0.99,
    n_papers: int = 8,
    raw_hits: int = 7,
    sensitivity_band: str = "stable",
    warnings: tuple[str, ...] = (),
) -> KPPosterior:
    return KPPosterior(
        kp_id=kp_id,
        n_papers=n_papers,
        raw_hits=raw_hits,
        weighted_hits=float(raw_hits),
        weighted_N=float(n_papers),
        lambda_used=0.2,
        tau_used=1.0,
        reference_year=2026,
        coverage_share=0.1,
        prior_alpha=0.1,
        prior_beta=0.9,
        posterior_alpha=7.1,
        posterior_beta=1.9,
        posterior_mean=posterior_mean,
        ci_lower_95=ci_lower_95,
        ci_upper_95=ci_upper_95,
        hotness_mean_share=0.1,
        hotness_std_share=0.02,
        trend_label="stable",
        trend_delta=0.0,
        trend_ci_95=(-0.1, 0.1),
        historical_mean=raw_hits / n_papers if n_papers else 0.0,
        tier=tier,
        tier_reasons=(f"posterior_mean={posterior_mean:.2f} >= 0.75",),
        sensitivity_band=sensitivity_band,
        warnings=warnings,
    )


def _coverage(
    kp_id: str,
    pattern_id: str,
    tier: str = "hot",
    raw_hits: int = 3,
    weighted_hits: float = 2.5,
    last_seen_year: float | None = 2024.4,
    saturation_index: float = 0.4,
    freshness_flag: bool = False,
    predicted_score: float = 0.85,
    occurrences: tuple[PatternOccurrence, ...] = (),
) -> PatternCoverage:
    return PatternCoverage(
        kp_id=kp_id,
        pattern_id=pattern_id,
        raw_hits=raw_hits,
        weighted_hits=weighted_hits,
        last_seen_year=last_seen_year,
        first_seen_year=2020.4 if last_seen_year else None,
        inter_arrival_years_mean=None,
        inter_arrival_years_max=None,
        saturation_index=saturation_index,
        freshness_flag=freshness_flag,
        predicted_score=predicted_score,
        complications_seen=(),
        complications_unseen=(),
        occurrences=occurrences,
        warnings=(),
        tier=tier,
        tier_reasons=(f"tier={tier}",),
    )


def _pattern_def(pattern_id: str, kp_id: str, label: str = "Test pattern") -> dict:
    return {
        "kp_id": kp_id,
        "pattern_id": pattern_id,
        "label": label,
        "solution_sketch": ["differentiate", "substitute named point"],
        "source": ["textbook §5.4 example 12"],
        "common_complications": [],
    }


def _question(year: float, qno: str, kp: str, pattern_id: str) -> dict:
    return {
        "year": year,
        "question_number": qno,
        "primary_kp": kp,
        "pattern_id": pattern_id,
        "prompt_summary": "Find tangent at the named point.",
        "complications": [],
    }


def _sweep(kp_id: str, distinct: tuple[str, ...] = ("anchor",)) -> SensitivitySweep:
    cell = SensitivityCell(
        lam=0.2,
        tau=1.0,
        posterior_mean=0.85,
        ci_lower_95=0.55,
        ci_upper_95=0.99,
        tier="anchor",
        warnings=(),
    )
    band = "stable" if len(distinct) <= 1 else "unstable"
    return SensitivitySweep(
        kp_id=kp_id,
        cells=(cell,),
        distinct_tiers=distinct,
        band=band,
    )


def _baseline_inputs():
    """Minimal but multi-KP input bundle."""
    posteriors = [
        _posterior(kp_id="L13.03", tier="anchor", posterior_mean=0.85, ci_lower_95=0.55),
        _posterior(
            kp_id="L14.01",
            tier="core",
            posterior_mean=0.55,
            ci_lower_95=0.30,
            ci_upper_95=0.78,
            n_papers=6,
            raw_hits=4,
        ),
        _posterior(
            kp_id="L01.01",
            tier="oneoff",
            posterior_mean=0.20,
            ci_lower_95=0.05,
            ci_upper_95=0.45,
            n_papers=5,
            raw_hits=1,
        ),
    ]
    sweeps = {p.kp_id: _sweep(p.kp_id) for p in posteriors}
    pattern_definitions = [
        _pattern_def("L13.03.P02", "L13.03", "Find tangent at the named point"),
        _pattern_def("L13.03.P05", "L13.03", "Vertical tangent edge case"),
        _pattern_def("L14.01.P01", "L14.01", "Implicit differentiation core"),
    ]
    occ = (PatternOccurrence(year=2024.4, question_number="7", confidence=1.0, is_primary=True),)
    pattern_coverage = [
        _coverage(
            kp_id="L13.03",
            pattern_id="L13.03.P02",
            tier="hot",
            raw_hits=3,
            weighted_hits=2.5,
            saturation_index=0.4,
            predicted_score=0.85,
            occurrences=occ,
        ),
        _coverage(
            kp_id="L13.03",
            pattern_id="L13.03.P05",
            tier="fresh",
            raw_hits=0,
            weighted_hits=0.0,
            last_seen_year=None,
            freshness_flag=True,
            predicted_score=0.30,
        ),
        _coverage(
            kp_id="L14.01",
            pattern_id="L14.01.P01",
            tier="hot",
            raw_hits=2,
            weighted_hits=1.6,
            predicted_score=0.65,
        ),
    ]
    mapping_questions = [
        _question(2024.4, "7", "L13.03", "L13.03.P02"),
    ]
    kps = [
        {"kp_id": "L13.03", "label": "Tangent / Normal", "lecture_prefix": "Lecture 13"},
        {"kp_id": "L14.01", "label": "Implicit differentiation"},
        {"kp_id": "L01.01", "label": "Number bases"},
    ]
    hyperparameters = {
        "course_id": "TEST-COURSE",
        "course_name": "Synthetic Course",
        "reference_year": 2026,
        "lambda": 0.2,
        "tau": 1.0,
        "alpha": 0.3,
        "lambda_grid": [0.0, 0.2, 0.4],
        "tau_grid": [0.5, 1.0, 2.0],
        "n_papers": 5,
        "n_kp": 3,
    }
    course_meta = {
        "course_id": "TEST-COURSE",
        "course_name": "Synthetic Course",
        "reference_year": 2026,
        "n_papers": 5,
        "n_kps": 3,
    }
    return {
        "posteriors": posteriors,
        "sweeps": sweeps,
        "pattern_definitions": pattern_definitions,
        "pattern_coverage": pattern_coverage,
        "mapping_questions": mapping_questions,
        "kps": kps,
        "hyperparameters": hyperparameters,
        "course_meta": course_meta,
    }


def _docx_headings(path: Path) -> list[str]:
    """Read every paragraph using a Heading style from a .docx file."""
    doc = Document(str(path))
    headings: list[str] = []
    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style else ""
        if style_name.startswith("Heading") or style_name == "Title":
            headings.append(para.text)
    return headings


def _docx_text(path: Path) -> str:
    """Concatenate every visible paragraph + table cell from a .docx file."""
    doc = Document(str(path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        chunks.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


# ---------------------------------------------------------------------------
# DOCX — happy-path
# ---------------------------------------------------------------------------


class TestWriteDocxEnglish:
    def test_creates_non_empty_file(self, tmp_path: Path) -> None:
        inputs = _baseline_inputs()
        out = tmp_path / "report.docx"
        result = write_docx(
            out,
            posteriors=inputs["posteriors"],
            sweeps=inputs["sweeps"],
            hyperparameters=inputs["hyperparameters"],
            pattern_coverage=inputs["pattern_coverage"],
            pattern_definitions=inputs["pattern_definitions"],
            mapping_questions=inputs["mapping_questions"],
            kps=inputs["kps"],
            course_meta=inputs["course_meta"],
            lang="en",
        )
        assert result == out
        assert out.exists()
        assert out.stat().st_size > 1000  # Comfortably non-empty.

    def test_section_headings_present(self, tmp_path: Path) -> None:
        inputs = _baseline_inputs()
        out = tmp_path / "report.docx"
        write_docx(
            out,
            posteriors=inputs["posteriors"],
            sweeps=inputs["sweeps"],
            hyperparameters=inputs["hyperparameters"],
            pattern_coverage=inputs["pattern_coverage"],
            pattern_definitions=inputs["pattern_definitions"],
            mapping_questions=inputs["mapping_questions"],
            kps=inputs["kps"],
            course_meta=inputs["course_meta"],
            lang="en",
        )
        headings = _docx_headings(out)
        joined = "\n".join(headings)
        assert SECTION_EXEC_SUMMARY in joined
        assert SECTION_PREDICTIONS in joined
        assert SECTION_CHEAT_SHEETS in joined
        assert SECTION_SENSITIVITY in joined
        assert APPENDIX_AUDIT in joined
        assert APPENDIX_PATTERNS in joined
        assert APPENDIX_METHODOLOGY in joined


# ---------------------------------------------------------------------------
# DOCX — pattern_definitions=None fallback
# ---------------------------------------------------------------------------


class TestWriteDocxNoPatternLayer:
    def test_renders_without_pattern_layer(self, tmp_path: Path) -> None:
        inputs = _baseline_inputs()
        out = tmp_path / "report-no-patterns.docx"
        write_docx(
            out,
            posteriors=inputs["posteriors"],
            sweeps=inputs["sweeps"],
            hyperparameters=inputs["hyperparameters"],
            pattern_coverage=None,
            pattern_definitions=None,
            mapping_questions=None,
            kps=None,
            course_meta=inputs["course_meta"],
            lang="en",
        )
        assert out.exists()
        assert out.stat().st_size > 500
        text = _docx_text(out)
        # Executive summary still lands.
        assert SECTION_EXEC_SUMMARY in text
        # Cheat-sheet section exists but contains the stub message.
        assert SECTION_CHEAT_SHEETS in text
        assert "No pattern data available" in text
        # Methodology appendix still ships.
        assert APPENDIX_METHODOLOGY in text


# ---------------------------------------------------------------------------
# Bilingual rendering
# ---------------------------------------------------------------------------


class TestWriteDocxBilingual:
    def test_lang_both_includes_zh_narrative_when_present(
        self, tmp_path: Path
    ) -> None:
        inputs = _baseline_inputs()
        narratives = {
            "L13.03": {
                "headline": "EN headline anchor",
                "narrative": "EN narrative paragraph.",
                "headline_zh": "ZH 标题：锚点必考。",
                "narrative_zh": "ZH 叙述：考切线。",
            },
        }
        out = tmp_path / "report-both.docx"
        write_docx(
            out,
            posteriors=inputs["posteriors"],
            sweeps=inputs["sweeps"],
            hyperparameters=inputs["hyperparameters"],
            pattern_coverage=inputs["pattern_coverage"],
            pattern_definitions=inputs["pattern_definitions"],
            mapping_questions=inputs["mapping_questions"],
            kps=inputs["kps"],
            tier_narratives=narratives,
            course_meta=inputs["course_meta"],
            lang="both",
        )
        text = _docx_text(out)
        assert "EN headline anchor" in text
        assert "ZH 标题：锚点必考。" in text
        assert "EN narrative paragraph." in text
        assert "ZH 叙述：考切线。" in text

    def test_lang_both_only_en_when_no_zh_narrative(
        self, tmp_path: Path
    ) -> None:
        inputs = _baseline_inputs()
        narratives = {
            "L13.03": {
                "headline": "EN-only headline",
                "narrative": "EN-only narrative paragraph.",
            },
        }
        out = tmp_path / "report-en-only.docx"
        write_docx(
            out,
            posteriors=inputs["posteriors"],
            sweeps=inputs["sweeps"],
            hyperparameters=inputs["hyperparameters"],
            pattern_coverage=inputs["pattern_coverage"],
            pattern_definitions=inputs["pattern_definitions"],
            mapping_questions=inputs["mapping_questions"],
            kps=inputs["kps"],
            tier_narratives=narratives,
            course_meta=inputs["course_meta"],
            lang="both",
        )
        text = _docx_text(out)
        assert "EN-only headline" in text
        assert "EN-only narrative paragraph." in text
        # No untranslated ZH leaks in when the input doesn't include any.
        assert "ZH 标题" not in text
        assert "ZH 叙述" not in text


# ---------------------------------------------------------------------------
# Appendix A — Full KP audit table contains every KP
# ---------------------------------------------------------------------------


class TestAppendixAuditTable:
    def test_contains_all_kp_ids(self, tmp_path: Path) -> None:
        inputs = _baseline_inputs()
        out = tmp_path / "report-audit.docx"
        write_docx(
            out,
            posteriors=inputs["posteriors"],
            sweeps=inputs["sweeps"],
            hyperparameters=inputs["hyperparameters"],
            pattern_coverage=inputs["pattern_coverage"],
            pattern_definitions=inputs["pattern_definitions"],
            mapping_questions=inputs["mapping_questions"],
            kps=inputs["kps"],
            course_meta=inputs["course_meta"],
            lang="en",
        )
        doc = Document(str(out))
        # Locate the table that follows the Appendix A heading.
        audit_table_kps: set[str] = set()
        for table in doc.tables:
            header_cells = [cell.text for cell in table.rows[0].cells]
            if header_cells and header_cells[0] == "kp_id" and "hotness_mean" in header_cells:
                for row in table.rows[1:]:
                    audit_table_kps.add(row.cells[0].text)
                break
        expected = {p.kp_id for p in inputs["posteriors"]}
        assert expected.issubset(audit_table_kps)


# ---------------------------------------------------------------------------
# Markdown parity
# ---------------------------------------------------------------------------


class TestMarkdownParity:
    def test_markdown_contains_expected_section_headers(
        self, tmp_path: Path
    ) -> None:
        inputs = _baseline_inputs()
        out = tmp_path / "report.md"
        result = write_markdown(
            out,
            inputs["posteriors"],
            inputs["sweeps"],
            inputs["hyperparameters"],
            pattern_coverage=inputs["pattern_coverage"],
            pattern_definitions=inputs["pattern_definitions"],
            mapping_questions=inputs["mapping_questions"],
            kps=inputs["kps"],
            course_meta=inputs["course_meta"],
        )
        assert result == out
        text = out.read_text()
        assert f"## {SECTION_EXEC_SUMMARY}" in text
        assert f"## {SECTION_PREDICTIONS}" in text
        assert f"## {SECTION_CHEAT_SHEETS}" in text
        assert f"## {APPENDIX_AUDIT}" in text
        assert f"## {APPENDIX_METHODOLOGY}" in text

    def test_markdown_audit_table_contains_every_kp(
        self, tmp_path: Path
    ) -> None:
        inputs = _baseline_inputs()
        out = tmp_path / "report.md"
        write_markdown(
            out,
            inputs["posteriors"],
            inputs["sweeps"],
            inputs["hyperparameters"],
            pattern_coverage=inputs["pattern_coverage"],
            pattern_definitions=inputs["pattern_definitions"],
            mapping_questions=inputs["mapping_questions"],
            kps=inputs["kps"],
            course_meta=inputs["course_meta"],
        )
        text = out.read_text()
        appendix_index = text.find(f"## {APPENDIX_AUDIT}")
        assert appendix_index >= 0
        section = text[appendix_index:]
        for posterior in inputs["posteriors"]:
            assert posterior.kp_id in section
