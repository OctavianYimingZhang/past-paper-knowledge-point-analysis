"""Render analysis outputs to Excel, JSON, Markdown, and Word.

This module is the only writer in the pipeline. All statistical work is
done by `statistical_model.py`, `sensitivity.py`, and `pattern_coverage.py`;
all extraction work is done by the `extract_*` modules. The writer takes
their dataclasses and emits files.

Design constraints:
- Every user-facing artifact must expose the hyperparameters used.
- Unstable KPs (sensitivity_band == "unstable") surface before the tier
  tables in the Markdown and DOCX summaries.
- No row is written without a warnings column, even if empty.
- The word "conjugate" never appears in output; the model is a
  moment-matched Beta posterior.
- The DOCX exporter is the canonical revision-plan deliverable. It always
  includes KP-frequency tables and sensitivity warnings; pattern-prediction
  sections are added when pattern coverage data is supplied.
- Default report language is English. Bilingual or Chinese-only output is
  opt-in via the `lang` parameter (`"en"`, `"zh"`, or `"both"`).
"""
from __future__ import annotations

import json
from dataclasses import asdict, is_dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Literal

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .sensitivity import (
    LeaveOneOutResult,
    SensitivitySweep,
    summarize_loo_for_report,
    summarize_sweep_for_report,
)
from .statistical_model import KPPosterior

ReportLang = Literal["en", "zh", "both"]


HEADER_FILL = PatternFill(start_color="FF305496", end_color="FF305496", fill_type="solid")
HEADER_FONT = Font(color="FFFFFFFF", bold=True)
WRAP = Alignment(wrap_text=True, vertical="top")


def _default(obj):
    if is_dataclass(obj):
        return asdict(obj)
    if isinstance(obj, set):
        return sorted(obj)
    if isinstance(obj, Path):
        return str(obj)
    raise TypeError(f"Unserializable type {type(obj).__name__}")


def write_json(
    out_path: str | Path,
    posteriors: list[KPPosterior],
    sweeps: dict[str, SensitivitySweep],
    loo: dict[str, LeaveOneOutResult],
    hyperparameters: dict[str, object],
) -> Path:
    """Write a single JSON payload with everything needed to reproduce the run."""
    path = Path(out_path)
    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "hyperparameters": hyperparameters,
        "posteriors": [asdict(p) for p in posteriors],
        "sensitivity_sweeps": {kp: asdict(s) for kp, s in sweeps.items()},
        "leave_one_out": {
            kp: {
                "baseline": asdict(r.baseline),
                "per_year": [[year, asdict(pos)] for year, pos in r.per_year],
                "max_abs_shift": r.max_abs_shift,
                "tier_flips": list(r.tier_flips),
            }
            for kp, r in loo.items()
        },
    }
    path.write_text(json.dumps(payload, indent=2, default=_default))
    return path


def write_excel(
    out_path: str | Path,
    posteriors: list[KPPosterior],
    sweeps: dict[str, SensitivitySweep],
    loo: dict[str, LeaveOneOutResult],
    hyperparameters: dict[str, object],
) -> Path:
    """Write the canonical Excel workbook."""
    path = Path(out_path)
    wb = Workbook()
    wb.remove(wb.active)

    _write_method_sheet(wb, hyperparameters)
    _write_predictions_sheet(wb, posteriors)
    _write_sensitivity_sheet(wb, sweeps)
    _write_loo_sheet(wb, loo)
    _write_trend_sheet(wb, posteriors)
    _write_review_sheet(wb, posteriors)

    wb.save(path)
    return path


def write_markdown(
    out_path: str | Path,
    posteriors: list[KPPosterior],
    sweeps: dict[str, SensitivitySweep],
    hyperparameters: dict[str, object],
) -> Path:
    """Write the executive-summary Markdown. Unstable KPs surface first."""
    path = Path(out_path)
    lines: list[str] = []
    lines.append(f"# Past-Paper Knowledge-Point Analysis")
    lines.append("")
    lines.append(f"_Generated: {datetime.now().isoformat(timespec='seconds')}_")
    lines.append("")
    lines.append("## Hyperparameters")
    lines.append("")
    for key, value in sorted(hyperparameters.items()):
        lines.append(f"- **{key}**: {value}")
    lines.append("")
    lines.append(
        "The posterior is a moment-matched Beta approximation under "
        "recency-weighted evidence. It is not a strict conjugate update."
    )
    lines.append("")

    unstable = [p for p in posteriors if p.sensitivity_band == "unstable"]
    if unstable:
        lines.append("## Unstable Results (read these first)")
        lines.append("")
        lines.append(
            "These knowledge points flipped tiers across the sensitivity sweep. "
            "Treat the headline tier with caution; see the `Sensitivity_Sweep` "
            "sheet for the full grid."
        )
        lines.append("")
        lines.append("| KP | Tier | P(mean) | CI95 | Distinct tiers in sweep |")
        lines.append("|----|------|---------|------|-------------------------|")
        for p in unstable:
            band = sweeps.get(p.kp_id)
            distinct = ", ".join(band.distinct_tiers) if band else "?"
            lines.append(
                f"| {p.kp_id} | {p.tier} | {p.posterior_mean:.2f} | "
                f"[{p.ci_lower_95:.2f}, {p.ci_upper_95:.2f}] | {distinct} |"
            )
        lines.append("")

    lines.append("## Tier Summary")
    lines.append("")
    tier_order = ("anchor", "core", "emerging", "legacy", "oneoff", "not_tested")
    for tier in tier_order:
        tier_rows = [p for p in posteriors if p.tier == tier]
        if not tier_rows:
            continue
        lines.append(f"### {tier.title()} ({len(tier_rows)} KPs)")
        lines.append("")
        lines.append("| KP | P(mean) | CI95 | Hotness | Trend | Warnings |")
        lines.append("|----|---------|------|---------|-------|----------|")
        for p in tier_rows:
            warn = "; ".join(p.warnings) if p.warnings else "-"
            lines.append(
                f"| {p.kp_id} | {p.posterior_mean:.2f} | "
                f"[{p.ci_lower_95:.2f}, {p.ci_upper_95:.2f}] | "
                f"{p.hotness_mean_share:.3f} +/- {p.hotness_std_share:.3f} | "
                f"{p.trend_label} | {warn} |"
            )
        lines.append("")

    path.write_text("\n".join(lines))
    return path


# ---------------------------------------------------------------------------
# Internals
# ---------------------------------------------------------------------------


def _write_method_sheet(wb: Workbook, hyperparameters: dict[str, object]) -> None:
    ws = wb.create_sheet("Method")
    rows: list[list[object]] = [
        ["Generated at", datetime.now().isoformat(timespec="seconds")],
        ["Posterior model", "Moment-matched Beta (not strict conjugate)"],
        ["Recency weighting", "w_i = exp(-lambda * (reference_year - year_i))"],
        ["Prior", "Beta(tau * coverage, tau * (1 - coverage)), tau in [0, 2]"],
        ["Trend", "Split-halves bootstrap of rate difference (no Mann-Kendall)"],
        ["Tier rules", "See references/tier-definitions.md"],
    ]
    for key, value in sorted(hyperparameters.items()):
        rows.append([key, _scalar(value)])
    _write_sheet(ws, ["Field", "Value"], rows)


def _scalar(value: object) -> object:
    """Coerce lists and other non-primitive values into an Excel-safe scalar."""
    if isinstance(value, (list, tuple, set)):
        return ", ".join(str(v) for v in value)
    if isinstance(value, dict):
        return json.dumps(value, sort_keys=True)
    return value


def _write_predictions_sheet(wb: Workbook, posteriors: list[KPPosterior]) -> None:
    ws = wb.create_sheet("Posterior_Predictions")
    headers = [
        "kp_id",
        "n_papers",
        "raw_hits",
        "weighted_hits",
        "weighted_N",
        "lambda_used",
        "tau_used",
        "coverage_share",
        "prior_alpha",
        "prior_beta",
        "posterior_alpha",
        "posterior_beta",
        "posterior_mean",
        "ci_lower_95",
        "ci_upper_95",
        "hotness_mean_share",
        "hotness_std_share",
        "trend_label",
        "trend_delta",
        "trend_ci_low",
        "trend_ci_high",
        "tier",
        "tier_reasons",
        "sensitivity_band",
        "warnings",
    ]
    rows = [
        [
            p.kp_id,
            p.n_papers,
            p.raw_hits,
            round(p.weighted_hits, 4),
            round(p.weighted_N, 4),
            p.lambda_used,
            p.tau_used,
            round(p.coverage_share, 4),
            round(p.prior_alpha, 4),
            round(p.prior_beta, 4),
            round(p.posterior_alpha, 4),
            round(p.posterior_beta, 4),
            round(p.posterior_mean, 4),
            round(p.ci_lower_95, 4),
            round(p.ci_upper_95, 4),
            round(p.hotness_mean_share, 4),
            round(p.hotness_std_share, 4),
            p.trend_label,
            round(p.trend_delta, 4),
            round(p.trend_ci_95[0], 4),
            round(p.trend_ci_95[1], 4),
            p.tier,
            " | ".join(p.tier_reasons),
            p.sensitivity_band,
            " | ".join(p.warnings),
        ]
        for p in posteriors
    ]
    _write_sheet(ws, headers, rows)


def _write_sensitivity_sheet(wb: Workbook, sweeps: dict[str, SensitivitySweep]) -> None:
    ws = wb.create_sheet("Sensitivity_Sweep")
    headers = [
        "kp_id",
        "lambda",
        "tau",
        "posterior_mean",
        "ci_lower_95",
        "ci_upper_95",
        "tier",
        "warnings",
        "band",
    ]
    rows: list[list[object]] = []
    for sweep in sweeps.values():
        for r in summarize_sweep_for_report(sweep):
            rows.append([r[h] for h in headers])
    _write_sheet(ws, headers, rows)


def _write_loo_sheet(wb: Workbook, loo: dict[str, LeaveOneOutResult]) -> None:
    ws = wb.create_sheet("Leave_One_Out")
    headers = [
        "kp_id",
        "dropped_year",
        "baseline_posterior_mean",
        "loo_posterior_mean",
        "shift",
        "abs_shift",
        "baseline_tier",
        "loo_tier",
        "tier_flipped",
    ]
    rows: list[list[object]] = []
    for result in loo.values():
        for r in summarize_loo_for_report(result):
            rows.append([r[h] for h in headers])
    _write_sheet(ws, headers, rows)


def _write_trend_sheet(wb: Workbook, posteriors: list[KPPosterior]) -> None:
    ws = wb.create_sheet("Trend_Analysis")
    headers = [
        "kp_id",
        "trend_label",
        "trend_delta",
        "trend_ci_low",
        "trend_ci_high",
        "historical_mean",
        "posterior_mean",
    ]
    rows = [
        [
            p.kp_id,
            p.trend_label,
            round(p.trend_delta, 4),
            round(p.trend_ci_95[0], 4),
            round(p.trend_ci_95[1], 4),
            round(p.historical_mean, 4),
            round(p.posterior_mean, 4),
        ]
        for p in posteriors
    ]
    _write_sheet(ws, headers, rows)


def _write_review_sheet(wb: Workbook, posteriors: list[KPPosterior]) -> None:
    ws = wb.create_sheet("Review_Queue")
    headers = ["kp_id", "tier", "warning"]
    rows: list[list[object]] = []
    for p in posteriors:
        for w in p.warnings:
            rows.append([p.kp_id, p.tier, w])
    _write_sheet(ws, headers, rows)


# ---------------------------------------------------------------------------
# DOCX exporter — pattern-aware revision-plan deliverable
# ---------------------------------------------------------------------------


def write_docx(
    out_path: str | Path,
    posteriors: list[KPPosterior],
    sweeps: dict[str, SensitivitySweep],
    hyperparameters: dict[str, object],
    pattern_coverage: dict[str, list[dict]] | None = None,
    pattern_definitions: dict[str, list[dict]] | None = None,
    tier_narratives: dict[str, dict] | None = None,
    lang: ReportLang = "en",
) -> Path:
    """Write the revision-plan Word document.

    Parameters
    ----------
    out_path
        Destination .docx file.
    posteriors
        KP-level Beta posteriors (existing pipeline).
    sweeps
        Sensitivity sweeps keyed by KP id.
    hyperparameters
        Run-level hyperparameters (lambda, tau, reference_year, etc.).
    pattern_coverage
        Optional. Map of `kp_id -> list[PatternCoverage-like dict]`. When
        present, the DOCX includes a "Pattern Predictions" section per
        anchor / core KP with already-tested vs still-possible breakdowns.
    pattern_definitions
        Optional. Map of `kp_id -> list[Pattern-like dict]` carrying the
        canonical taxonomy (label, asked_operation, solution_sketch, etc.).
        Required if `pattern_coverage` is supplied.
    tier_narratives
        Optional. Map of `kp_id -> {"narrative": str, ...}` from the Opus
        statistical-interpreter stage. Used to seed per-KP commentary.
    lang
        Report language. "en" (default), "zh", or "both".
    """
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover - guarded by requirements.txt
        raise RuntimeError(
            "python-docx is required for DOCX output; pip install python-docx"
        ) from exc

    path = Path(out_path)
    doc = Document()
    _set_default_font(doc)

    course_name = str(hyperparameters.get("course_name", "Past-Paper Analysis"))
    course_id = str(hyperparameters.get("course_id", "course"))
    n_papers = hyperparameters.get("n_papers", "?")
    n_kp = hyperparameters.get("n_kp", "?")
    n_patterns = sum(len(v) for v in (pattern_definitions or {}).values()) if pattern_definitions else 0

    # ----- Cover page -----
    doc.add_heading(course_name, level=0)
    cover = doc.add_paragraph()
    cover.add_run(_t("Past-Paper Knowledge-Point & Pattern Analysis", "历年试卷知识点与考法分析", lang)).bold = True
    doc.add_paragraph(
        _t(
            f"Course code: {course_id}",
            f"课程代码：{course_id}",
            lang,
        )
    )
    doc.add_paragraph(
        _t(
            f"Generated: {datetime.now().isoformat(timespec='seconds')}",
            f"生成时间：{datetime.now().isoformat(timespec='seconds')}",
            lang,
        )
    )
    doc.add_paragraph(
        _t(
            f"Papers analysed: {n_papers}    "
            f"Knowledge points: {n_kp}    "
            f"Patterns: {n_patterns if n_patterns else 'n/a (KP-only run)'}",
            f"分析卷数：{n_papers}    知识点：{n_kp}    "
            f"考法 / patterns：{n_patterns if n_patterns else '本次未启用'}",
            lang,
        )
    )

    # ----- Methodology -----
    doc.add_heading(_t("Methodology", "方法说明", lang), level=1)
    method_lines = [
        _t(
            "The KP-frequency layer uses a moment-matched Beta posterior over "
            "recency-weighted hits with a curriculum-coverage prior. The pattern "
            "layer uses transparent frequency + recency statistics (no Beta "
            "posterior) because per-pattern data is too sparse to support a "
            "credible interval honestly.",
            "知识点层使用矩匹配 Beta 后验，结合时间衰减权重与课本覆盖先验；"
            "考法层使用透明的频率 + 时间衰减统计，不使用 Beta 后验，原因是"
            "单考法的数据稀疏，无法诚实地给出可信区间。",
            lang,
        ),
        _t(
            f"Hyperparameters: lambda={hyperparameters.get('lambda')}, "
            f"tau={hyperparameters.get('tau')}, "
            f"reference_year={hyperparameters.get('reference_year')}, "
            f"alpha (novelty bias)={hyperparameters.get('alpha', 0.3)}.",
            f"超参数：lambda={hyperparameters.get('lambda')}, "
            f"tau={hyperparameters.get('tau')}, "
            f"参考年份={hyperparameters.get('reference_year')}, "
            f"alpha（新颖性偏置）={hyperparameters.get('alpha', 0.3)}。",
            lang,
        ),
    ]
    for line in method_lines:
        doc.add_paragraph(line)

    # ----- Unstable KPs first -----
    unstable = [p for p in posteriors if p.sensitivity_band == "unstable"]
    if unstable:
        doc.add_heading(_t("Unstable KPs (read first)", "敏感度不稳定（请先阅读）", lang), level=1)
        doc.add_paragraph(
            _t(
                "These knowledge points flipped tiers across the (lambda, tau) "
                "sweep. Treat the headline tier with caution.",
                "下列知识点在 (lambda, tau) 扫描中档位发生过切换，标题档位需谨慎对待。",
                lang,
            )
        )
        _add_kp_table(
            doc,
            unstable,
            sweeps,
            include_distinct_tiers=True,
            lang=lang,
        )

    # ----- Section A: KP frequency tier tables -----
    doc.add_heading(_t("A. Knowledge-point frequency", "A. 知识点出现频率", lang), level=1)
    tier_order = ("anchor", "core", "emerging", "legacy", "oneoff", "not_tested")
    tier_titles_en = {
        "anchor": "Anchor — almost certain",
        "core": "Core — high probability",
        "emerging": "Emerging — rising",
        "legacy": "Legacy — cooling",
        "oneoff": "One-off / sparse",
        "not_tested": "Not tested in mapped years",
    }
    tier_titles_zh = {
        "anchor": "锚点 — 几乎必考",
        "core": "核心 — 高概率",
        "emerging": "上升题",
        "legacy": "退潮题",
        "oneoff": "偶现题",
        "not_tested": "未考查",
    }
    for tier in tier_order:
        tier_rows = [p for p in posteriors if p.tier == tier]
        if not tier_rows:
            continue
        title = _t(tier_titles_en[tier], tier_titles_zh[tier], lang)
        doc.add_heading(f"{title} ({len(tier_rows)})", level=2)
        _add_kp_table(doc, tier_rows, sweeps, include_distinct_tiers=False, lang=lang)

    # ----- Section B: Pattern predictions per anchor / core KP -----
    if pattern_coverage and pattern_definitions:
        doc.add_heading(
            _t("B. Pattern predictions (test methods)", "B. 考法预测", lang),
            level=1,
        )
        doc.add_paragraph(
            _t(
                "For each anchor / core KP, the following sub-sections decompose "
                "the recurrence into question patterns ('how it has been tested'). "
                "The 'Already tested' table cites specific (year, question) "
                "appearances per pattern. The 'Still possible' table flags "
                "patterns that the textbook or lectures cover but the examiner "
                "has not used recently.",
                "对每个锚点 / 核心知识点，下列小节将其复发分解为不同的"
                "考法（题型）。'已经怎么考过了'表给出每个考法对应的具体年份与题号；"
                "'还能怎么考'表标记课本或讲义中存在但近年未出的考法。",
                lang,
            )
        )
        priority_kps = [p for p in posteriors if p.tier in ("anchor", "core")]
        for posterior in priority_kps:
            patterns = pattern_definitions.get(posterior.kp_id, [])
            coverage = pattern_coverage.get(posterior.kp_id, [])
            if not patterns:
                continue
            heading = posterior.kp_id
            doc.add_heading(heading, level=2)
            narrative = (tier_narratives or {}).get(posterior.kp_id, {})
            if narrative.get("narrative"):
                doc.add_paragraph(narrative["narrative"])
            _add_pattern_decomposition_table(doc, patterns, coverage, lang=lang)
            _add_already_tested_table(doc, coverage, lang=lang)
            _add_still_possible_table(doc, patterns, coverage, lang=lang)
            _add_solution_sketch(doc, patterns, coverage, lang=lang)

    # ----- Section C: Sensitivity, leave-one-out, warnings -----
    doc.add_heading(
        _t("C. Sensitivity & data warnings", "C. 敏感度与数据警告", lang), level=1
    )
    if sweeps:
        doc.add_paragraph(
            _t(
                f"Sensitivity sweep grid: lambda in {hyperparameters.get('lambda_grid')}, "
                f"tau in {hyperparameters.get('tau_grid')}.",
                f"敏感度扫描网格：lambda ∈ {hyperparameters.get('lambda_grid')}, "
                f"tau ∈ {hyperparameters.get('tau_grid')}。",
                lang,
            )
        )
    warned = [p for p in posteriors if p.warnings]
    if warned:
        doc.add_heading(
            _t("Data warnings per KP", "知识点级数据警告", lang), level=2
        )
        for p in warned:
            line = doc.add_paragraph(style="List Bullet")
            line.add_run(f"{p.kp_id} — {p.tier}: ").bold = True
            line.add_run("; ".join(p.warnings))

    # ----- Appendix: full pattern catalogue -----
    if pattern_definitions:
        doc.add_heading(
            _t("Appendix — pattern catalogue", "附录 — 考法目录", lang), level=1
        )
        for kp_id, patterns in sorted(pattern_definitions.items()):
            doc.add_heading(kp_id, level=3)
            for p in patterns:
                pid = p.get("pattern_id", "?")
                label = p.get("label", "")
                line = doc.add_paragraph(style="List Bullet")
                line.add_run(f"{pid} — ").bold = True
                line.add_run(label)
                src = p.get("source") or []
                if src:
                    src_para = doc.add_paragraph()
                    src_para.add_run(
                        _t("    Source: ", "    出处：", lang)
                    ).italic = True
                    src_para.add_run("; ".join(src))

    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------


def _set_default_font(doc) -> None:  # pragma: no cover - cosmetic
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _t(en: str, zh: str, lang: ReportLang) -> str:
    """Render text in the requested language."""
    if lang == "en":
        return en
    if lang == "zh":
        return zh
    return f"{en}\n{zh}"


def _add_kp_table(
    doc,
    posteriors: list[KPPosterior],
    sweeps: dict[str, SensitivitySweep],
    include_distinct_tiers: bool,
    lang: ReportLang,
) -> None:
    headers_en = ["KP", "Tier", "P(mean)", "CI95", "Hits", "Trend", "Warnings"]
    headers_zh = ["KP", "档位", "P(均值)", "CI95", "命中", "趋势", "警告"]
    if include_distinct_tiers:
        headers_en.append("Sweep tiers")
        headers_zh.append("扫描档位")
    headers = [_t(en, zh, lang) for en, zh in zip(headers_en, headers_zh)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for p in posteriors:
        row = table.add_row().cells
        row[0].text = p.kp_id
        row[1].text = p.tier
        row[2].text = f"{p.posterior_mean:.2f}"
        row[3].text = f"[{p.ci_lower_95:.2f}, {p.ci_upper_95:.2f}]"
        row[4].text = f"{p.raw_hits}/{p.n_papers}"
        row[5].text = p.trend_label
        row[6].text = "; ".join(p.warnings) if p.warnings else "-"
        if include_distinct_tiers:
            sweep = sweeps.get(p.kp_id)
            row[7].text = ", ".join(sweep.distinct_tiers) if sweep else "?"


def _add_pattern_decomposition_table(
    doc,
    patterns: list[dict],
    coverage: list[dict],
    lang: ReportLang,
) -> None:
    doc.add_paragraph(
        _t("Pattern decomposition:", "考法分解：", lang)
    ).runs[0].bold = True
    cov_by_pid = {c["pattern_id"]: c for c in coverage}
    headers = [
        _t("Pattern", "考法", lang),
        _t("Label", "标签", lang),
        _t("Hits", "命中", lang),
        _t("Last seen", "最近一次", lang),
        _t("Saturation", "饱和度", lang),
        _t("Fresh", "未考过", lang),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for p in patterns:
        pid = p.get("pattern_id", "?")
        cov = cov_by_pid.get(pid, {})
        row = table.add_row().cells
        row[0].text = pid
        row[1].text = p.get("label", "")
        row[2].text = str(cov.get("raw_hits", 0))
        row[3].text = str(cov.get("last_seen_year") or "—")
        sat = cov.get("saturation_index")
        row[4].text = f"{sat:.2f}" if isinstance(sat, (int, float)) else "—"
        row[5].text = (
            _t("yes", "是", lang) if cov.get("freshness_flag") else _t("no", "否", lang)
        )


def _add_already_tested_table(
    doc, coverage: list[dict], lang: ReportLang
) -> None:
    doc.add_paragraph(
        _t("Already tested (year × question):", "已经怎么考过了（年份 × 题号）：", lang)
    ).runs[0].bold = True
    rows = []
    for cov in coverage:
        for occ in cov.get("occurrences", []):
            rows.append(
                {
                    "pattern_id": cov.get("pattern_id"),
                    "year": occ.get("year"),
                    "question": occ.get("question_number"),
                    "complications": ", ".join(occ.get("complications") or []) or "—",
                }
            )
    if not rows:
        doc.add_paragraph(
            _t("(no recorded occurrences)", "（暂无记录）", lang)
        )
        return
    headers = [
        _t("Pattern", "考法", lang),
        _t("Year", "年份", lang),
        _t("Question", "题号", lang),
        _t("Complications used", "出现的扩展", lang),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        row = table.add_row().cells
        row[0].text = str(r["pattern_id"])
        row[1].text = str(r["year"])
        row[2].text = str(r["question"])
        row[3].text = r["complications"]


def _add_still_possible_table(
    doc,
    patterns: list[dict],
    coverage: list[dict],
    lang: ReportLang,
) -> None:
    doc.add_paragraph(
        _t(
            "Still possible (textbook/lecture-seeded but unseen recently):",
            "还能怎么考（课本/讲义涵盖但近年未出）：",
            lang,
        )
    ).runs[0].bold = True
    cov_by_pid = {c["pattern_id"]: c for c in coverage}
    rows: list[dict] = []
    for p in patterns:
        pid = p.get("pattern_id", "?")
        cov = cov_by_pid.get(pid, {})
        if cov.get("freshness_flag"):
            rows.append(
                {
                    "pattern_id": pid,
                    "label": p.get("label", ""),
                    "complications": ", ".join(
                        cov.get("complications_unseen") or p.get("common_complications") or []
                    )
                    or "—",
                }
            )
        elif cov.get("complications_unseen"):
            rows.append(
                {
                    "pattern_id": pid,
                    "label": f"{p.get('label', '')}  "
                    + _t("(seen, but with new complications)", "（已考但有新扩展）", lang),
                    "complications": ", ".join(cov["complications_unseen"]),
                }
            )
    if not rows:
        doc.add_paragraph(
            _t(
                "(no fresh patterns flagged for this KP)",
                "（该知识点暂无未出现的新考法）",
                lang,
            )
        )
        return
    headers = [
        _t("Pattern", "考法", lang),
        _t("Label", "标签", lang),
        _t("Suggested complications", "建议的扩展", lang),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        row = table.add_row().cells
        row[0].text = str(r["pattern_id"])
        row[1].text = r["label"]
        row[2].text = r["complications"]


def _add_solution_sketch(
    doc,
    patterns: list[dict],
    coverage: list[dict],
    lang: ReportLang,
) -> None:
    """Pick the most-likely pattern for the KP and emit its solution sketch."""
    cov_by_pid = {c["pattern_id"]: c for c in coverage}
    ranked = sorted(
        patterns,
        key=lambda p: cov_by_pid.get(p.get("pattern_id", ""), {}).get("predicted_score", 0.0),
        reverse=True,
    )
    if not ranked:
        return
    leader = ranked[0]
    sketch = leader.get("solution_sketch") or []
    if not sketch:
        return
    doc.add_paragraph(
        _t(
            f"Solution sketch — most-likely pattern ({leader.get('pattern_id', '?')}):",
            f"解题思路 — 最可能考法（{leader.get('pattern_id', '?')}）：",
            lang,
        )
    ).runs[0].bold = True
    for step in sketch:
        line = doc.add_paragraph(style="List Number")
        line.add_run(str(step))


def _write_sheet(ws, headers: list[str], rows: Iterable[list[object]]) -> None:
    ws.append(headers)
    for cell in ws[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = WRAP
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"
    for row in rows:
        ws.append(list(row))
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = WRAP
    for column_cells in ws.columns:
        width = 14
        for cell in column_cells:
            if cell.value is None:
                continue
            width = max(width, min(len(str(cell.value)) + 2, 60))
        ws.column_dimensions[get_column_letter(column_cells[0].column)].width = width
