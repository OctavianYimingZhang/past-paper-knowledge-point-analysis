"""DOCX writer with executive summary, per-KP cheat-sheets, and appendices.

Layout (in order):

1. Executive Summary (≤1 page) — headline, top focus KPs table, top fresh
   patterns table, biggest caveats bullets.
2. KP Predictions — narrow 5-column scannable table.
3. How It Will Be Tested — per-KP cheat-sheets for anchor + core KPs only.
4. Sensitivity & Warnings — unstable KPs, leave-one-out shifts, warnings.
5. Appendix A — Full KP Audit Table.
6. Appendix B — Pattern Catalogue.
7. Appendix C — Methodology.

Every heading uses python-docx's "Heading 1/2/3" styles so the document
ships a working table-of-contents structure. Tables apply the
"Light Grid Accent 1" style when available and fall back to "Table Grid".
"""
from __future__ import annotations

from pathlib import Path

from ..pattern_coverage import PatternCoverage
from ..sensitivity import LeaveOneOutResult, SensitivitySweep
from ..statistical_model import KPPosterior
from ._common import (
    APPENDIX_AUDIT,
    APPENDIX_METHODOLOGY,
    APPENDIX_PATTERNS,
    KPCheatSheet,
    MAX_ALREADY_TESTED_ROWS,
    MAX_STILL_POSSIBLE_ROWS,
    ReportLang,
    SECTION_CHEAT_SHEETS,
    SECTION_EXEC_SUMMARY,
    SECTION_PREDICTIONS,
    SECTION_SENSITIVITY,
    build_sheets,
    caveat_bullets,
    confidence_chip,
    confidence_chip_zh,
    executive_headline,
    label_index,
    render_text,
    resolve_course_meta,
    scalar_str,
    sort_for_predictions_table,
    top_focus_kps,
    top_fresh_targets,
)


def write_docx(
    out_path: str | Path,
    *,
    posteriors: list[KPPosterior],
    sweeps: dict[str, SensitivitySweep] | list,
    hyperparameters: dict,
    pattern_coverage: list[PatternCoverage] | None = None,
    pattern_definitions: list[dict] | None = None,
    mapping_questions: list[dict] | None = None,
    kps: list[dict] | None = None,
    tier_narratives: dict | None = None,
    course_meta: dict | None = None,
    loo: dict[str, LeaveOneOutResult] | None = None,
    lang: ReportLang = "en",
) -> Path:
    """Write the revision-plan Word document with the new layout."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - guarded by requirements.txt
        raise RuntimeError(
            "python-docx is required for DOCX output; pip install python-docx"
        ) from exc

    path = Path(out_path)
    sweeps_dict = _normalize_sweeps(sweeps)
    sheets = build_sheets(
        posteriors=posteriors,
        pattern_coverage=pattern_coverage,
        pattern_definitions=pattern_definitions,
        mapping_questions=mapping_questions,
        kps=kps,
        tier_narratives=tier_narratives,
    )
    course = resolve_course_meta(course_meta, hyperparameters, posteriors)

    doc = Document()
    _set_default_font(doc)

    _render_executive_summary(
        doc=doc,
        posteriors=posteriors,
        sheets=sheets,
        course_meta=course,
        lang=lang,
    )
    _add_page_break(doc)

    _render_predictions_table(
        doc=doc, posteriors=posteriors, sheets=sheets, lang=lang
    )
    _add_page_break(doc)

    _render_kp_cheatsheets(
        doc=doc,
        posteriors=posteriors,
        sheets=sheets,
        tier_narratives=tier_narratives or {},
        lang=lang,
    )

    _render_sensitivity_section(
        doc=doc,
        posteriors=posteriors,
        sweeps=sweeps_dict,
        loo=loo or {},
        lang=lang,
    )
    _add_page_break(doc)

    _render_full_kp_audit(doc=doc, posteriors=posteriors, lang=lang)
    _add_page_break(doc)

    _render_pattern_catalogue(
        doc=doc,
        pattern_coverage=pattern_coverage,
        pattern_definitions=pattern_definitions,
        lang=lang,
    )
    _add_page_break(doc)

    _render_methodology_appendix(
        doc=doc,
        hyperparameters=hyperparameters,
        lang=lang,
    )

    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Section helpers
# ---------------------------------------------------------------------------


def _render_executive_summary(
    doc,
    posteriors: list[KPPosterior],
    sheets: dict[str, KPCheatSheet] | None,
    course_meta: dict,
    lang: ReportLang,
) -> None:
    course_name = course_meta.get("course_name", "Past-Paper Analysis")
    course_id = course_meta.get("course_id", "course")
    reference_year = course_meta.get("reference_year", "?")
    n_papers = course_meta.get("n_papers", "?")
    n_kps = course_meta.get("n_kps", len(posteriors))

    doc.add_heading(course_name, level=0)
    doc.add_heading(render_text(SECTION_EXEC_SUMMARY, "执行摘要", lang), level=1)
    meta = doc.add_paragraph()
    meta.add_run(
        render_text(
            f"Course {course_id} • Reference year {reference_year} • "
            f"{n_papers} formal paper(s) • {n_kps} knowledge points",
            f"课程 {course_id} • 参考年份 {reference_year} • "
            f"{n_papers} 份正式试卷 • {n_kps} 个知识点",
            lang,
        )
    ).italic = True

    doc.add_paragraph(executive_headline(posteriors, course_meta, sheets))

    doc.add_heading(render_text("Top focus KPs", "重点知识点", lang), level=2)
    headers_en = ["KP", "Tier", "Posterior", "Headline pattern"]
    headers_zh = ["知识点", "档位", "后验", "主考法"]
    headers = [render_text(en, zh, lang) for en, zh in zip(headers_en, headers_zh)]
    focus = top_focus_kps(posteriors, sheets)
    rows: list[list[str]] = []
    if not focus:
        rows.append(
            [render_text("(none detected)", "（暂无）", lang), "", "", ""]
        )
    else:
        for posterior, sheet in focus:
            label = sheet.kp_label if sheet else posterior.kp_id
            headline_pattern = "—"
            if sheet and sheet.dominant_pattern is not None:
                dom = sheet.dominant_pattern
                headline_pattern = f"{dom.pattern_id} — {dom.pattern_label}"
            posterior_text = (
                f"{posterior.posterior_mean:.2f} "
                f"[{posterior.ci_lower_95:.2f}, {posterior.ci_upper_95:.2f}]"
            )
            rows.append(
                [
                    f"{posterior.kp_id} — {label}",
                    posterior.tier,
                    posterior_text,
                    headline_pattern,
                ]
            )
    _add_table_with_headers(doc, headers, rows)

    doc.add_heading(
        render_text("Top fresh patterns to drill", "重点未考考法", lang), level=2
    )
    headers_en = ["KP", "Pattern", "Why fresh", "Source"]
    headers_zh = ["知识点", "考法", "未考原因", "出处"]
    headers = [render_text(en, zh, lang) for en, zh in zip(headers_en, headers_zh)]
    fresh_rows: list[list[str]] = []
    fresh = top_fresh_targets(sheets)
    if not fresh:
        fresh_rows.append(
            [
                render_text("(no fresh patterns flagged)", "（暂无未考考法）", lang),
                "",
                "",
                "",
            ]
        )
    else:
        for sheet, variant in fresh:
            sources = "; ".join(variant.sources) if variant.sources else "—"
            fresh_rows.append(
                [
                    sheet.kp_id,
                    f"{variant.pattern_id} — {variant.pattern_label}",
                    variant.rationale,
                    sources,
                ]
            )
    _add_table_with_headers(doc, headers, fresh_rows)

    bullets = caveat_bullets(posteriors)
    if bullets:
        doc.add_heading(render_text("Biggest caveats", "主要警告", lang), level=2)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")


def _render_predictions_table(
    doc,
    posteriors: list[KPPosterior],
    sheets: dict[str, KPCheatSheet] | None,
    lang: ReportLang,
) -> None:
    doc.add_heading(render_text(SECTION_PREDICTIONS, "知识点预测", lang), level=1)
    doc.add_paragraph(
        render_text(
            "Sorted by tier, then posterior_mean descending. The 'Confidence' "
            "chip is high when ci_lower_95 ≥ 0.50, medium when posterior_mean ≥ "
            "0.40, low otherwise.",
            "按档位、后验均值排序。'置信度'：当 ci_lower_95 ≥ 0.50 为高，"
            "posterior_mean ≥ 0.40 为中，其余为低。",
            lang,
        )
    )
    headers_en = ["KP ID", "Label", "Hits / Papers", "Tier", "Confidence"]
    headers_zh = ["知识点 ID", "标签", "命中 / 卷数", "档位", "置信度"]
    headers = [render_text(en, zh, lang) for en, zh in zip(headers_en, headers_zh)]
    rows: list[list[str]] = []
    for posterior in sort_for_predictions_table(posteriors):
        chip = confidence_chip(posterior)
        chip_label = render_text(chip, confidence_chip_zh(chip), lang)
        sheet = sheets.get(posterior.kp_id) if sheets else None
        label = sheet.kp_label if sheet and sheet.kp_label else posterior.kp_id
        rows.append(
            [
                posterior.kp_id,
                label,
                f"{posterior.raw_hits}/{posterior.n_papers}",
                posterior.tier,
                chip_label,
            ]
        )
    _add_table_with_headers(doc, headers, rows)


def _render_kp_cheatsheets(
    doc,
    posteriors: list[KPPosterior],
    sheets: dict[str, KPCheatSheet] | None,
    tier_narratives: dict,
    lang: ReportLang,
) -> None:
    doc.add_heading(render_text(SECTION_CHEAT_SHEETS, "考法详解", lang), level=1)
    if sheets is None:
        doc.add_paragraph(
            render_text(
                "No pattern data available for this run. Run extract-textbook + "
                "pattern-architect + pattern-classifier + pattern-coverage to "
                "enable per-KP cheat-sheets.",
                "本次未提供考法数据。运行 extract-textbook + pattern-architect + "
                "pattern-classifier + pattern-coverage 以启用考法详解。",
                lang,
            )
        )
        return
    primary = [p for p in posteriors if p.tier in ("anchor", "core")]
    primary.sort(key=lambda p: (-p.posterior_mean, p.kp_id))
    if primary:
        focus = primary
    else:
        fallback = [
            p for p in posteriors if p.tier in ("emerging", "oneoff", "legacy")
        ]
        fallback.sort(key=lambda p: (-p.posterior_mean, p.kp_id))
        focus = fallback
        if focus:
            doc.add_paragraph(
                render_text(
                    "No anchor or core KPs detected (data is too thin to clear "
                    "those tiers). Cheat-sheets below cover the highest-posterior "
                    "KPs that did appear in the available papers — treat them as "
                    "drill targets pending more papers.",
                    "数据稀疏，未触发锚点或核心层；以下提供出现频次最高的 KP 考法详解，"
                    "可作为优先复习目标。",
                    lang,
                )
            )
    if not focus:
        doc.add_paragraph(
            render_text(
                "No KPs with exam evidence; only curriculum-only inferences.",
                "未检测到任何题目对应的知识点；仅有课程材料推断。",
                lang,
            )
        )
        return
    for posterior in focus:
        sheet = sheets.get(posterior.kp_id)
        if sheet is None:
            continue
        narrative = tier_narratives.get(posterior.kp_id) or {}
        _render_kp_cheatsheet(doc, sheet, narrative, lang)
        _add_page_break(doc)


def _render_kp_cheatsheet(
    doc,
    sheet: KPCheatSheet,
    narrative: dict,
    lang: ReportLang,
) -> None:
    """Render one KP cheat-sheet card."""
    title = f"{sheet.kp_id} — {sheet.kp_label}  [{sheet.tier}]"
    doc.add_heading(title, level=2)
    if sheet.lecture_ref:
        sub = doc.add_paragraph()
        sub.add_run(f"({sheet.lecture_ref})").italic = True
    if sheet.past_paper_refs:
        refs = doc.add_paragraph()
        refs.add_run(render_text("Past papers: ", "历年试卷：", lang)).bold = True
        refs.add_run("; ".join(sheet.past_paper_refs))
    posterior_para = doc.add_paragraph()
    posterior_para.add_run(render_text("Posterior: ", "后验：", lang)).bold = True
    posterior_para.add_run(sheet.posterior_summary)

    headline_en = sheet.headline
    headline_zh = str(narrative.get("headline_zh", "")).strip() if narrative else ""
    narrative_en = sheet.narrative
    narrative_zh = str(narrative.get("narrative_zh", "")).strip() if narrative else ""

    if lang == "en":
        if headline_en:
            doc.add_paragraph(headline_en)
        if narrative_en:
            doc.add_paragraph(narrative_en)
    elif lang == "zh":
        if headline_zh:
            doc.add_paragraph(headline_zh)
        elif headline_en:
            doc.add_paragraph(headline_en)
        if narrative_zh:
            doc.add_paragraph(narrative_zh)
        elif narrative_en:
            doc.add_paragraph(narrative_en)
    else:  # both
        if headline_en:
            doc.add_paragraph(headline_en)
        if headline_zh:
            doc.add_paragraph(headline_zh)
        if narrative_en:
            doc.add_paragraph(narrative_en)
        if narrative_zh:
            doc.add_paragraph(narrative_zh)

    # H3 — How it will be tested
    doc.add_heading(render_text("How it will be tested", "考法预测", lang), level=3)
    if sheet.dominant_pattern is not None:
        dom = sheet.dominant_pattern
        para = doc.add_paragraph()
        para.add_run(render_text("Dominant pattern: ", "主考法：", lang)).bold = True
        para.add_run(f"{dom.pattern_id} — {dom.pattern_label}. ")
        para.add_run(dom.rationale + ".")
        if dom.solution_sketch:
            for step in dom.solution_sketch:
                doc.add_paragraph(str(step), style="List Bullet")
    else:
        doc.add_paragraph(
            render_text(
                "No dominant pattern derivable for this KP.",
                "该知识点无法推导主考法。",
                lang,
            )
        )
    context = doc.add_paragraph()
    context.add_run(render_text("Saturated: ", "饱和：", lang)).bold = True
    context.add_run(
        ", ".join(v.pattern_id for v in sheet.saturated_patterns) or "—"
    )
    context.add_run(render_text("  Fresh: ", "  未考：", lang)).bold = True
    context.add_run(", ".join(v.pattern_id for v in sheet.fresh_patterns) or "—")
    context.add_run(render_text("  Dormant: ", "  休眠：", lang)).bold = True
    context.add_run(", ".join(v.pattern_id for v in sheet.dormant_patterns) or "—")

    # H3 — Already tested
    doc.add_heading(render_text("Already tested", "已经怎么考过", lang), level=3)
    if sheet.already_tested:
        headers = [
            render_text("Year/Sitting", "年份", lang),
            render_text("Q", "题号", lang),
            render_text("Pattern", "考法", lang),
            render_text("Prompt summary", "题干摘要", lang),
        ]
        rows: list[list[str]] = []
        for ex in sheet.already_tested[:MAX_ALREADY_TESTED_ROWS]:
            rows.append(
                [
                    ex.year_label,
                    f"Q{ex.question_number}",
                    ex.pattern_id,
                    ex.prompt_summary or "—",
                ]
            )
        _add_table_with_headers(doc, headers, rows)
    else:
        doc.add_paragraph(
            render_text("No recorded occurrences.", "暂无记录。", lang)
        )

    # H3 — Still possible
    doc.add_heading(render_text("Still possible", "还能怎么考", lang), level=3)
    if sheet.still_possible:
        headers = [
            render_text("Pattern", "考法", lang),
            render_text("Why it's possible", "为何可能", lang),
            render_text("Source", "出处", lang),
        ]
        rows = []
        for variant in sheet.still_possible[:MAX_STILL_POSSIBLE_ROWS]:
            sources = "; ".join(variant.sources) if variant.sources else "—"
            rows.append(
                [
                    f"{variant.pattern_id} — {variant.pattern_label}",
                    variant.rationale,
                    sources,
                ]
            )
        _add_table_with_headers(doc, headers, rows)
    else:
        doc.add_paragraph(
            render_text(
                "No fresh or saturated-with-unseen patterns flagged.",
                "暂无未考或可扩展的考法。",
                lang,
            )
        )

    # H3 — Drill set
    doc.add_heading(render_text("Drill set", "练习清单", lang), level=3)
    if sheet.drill_set:
        for item in sheet.drill_set:
            doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph(
            render_text("No drill set assembled.", "暂无练习清单。", lang)
        )

    if sheet.open_caveats:
        caveat_para = doc.add_paragraph()
        caveat_para.add_run(render_text("Caveats: ", "警告：", lang)).bold = True
        caveat_para.add_run(" | ".join(sheet.open_caveats))


def _render_sensitivity_section(
    doc,
    posteriors: list[KPPosterior],
    sweeps: dict[str, SensitivitySweep],
    loo: dict[str, LeaveOneOutResult],
    lang: ReportLang,
) -> None:
    doc.add_heading(render_text(SECTION_SENSITIVITY, "敏感度与警告", lang), level=1)
    unstable = [p for p in posteriors if p.sensitivity_band == "unstable"]
    if unstable:
        doc.add_heading(render_text("Unstable KPs", "不稳定知识点", lang), level=2)
        headers = [
            render_text("KP ID", "知识点 ID", lang),
            render_text("Label", "标签", lang),
            render_text("Sensitivity band", "敏感度", lang),
            render_text("Distinct sweep tiers", "扫描档位", lang),
        ]
        rows: list[list[str]] = []
        for p in unstable:
            sweep = sweeps.get(p.kp_id)
            distinct = ", ".join(sweep.distinct_tiers) if sweep else "?"
            rows.append([p.kp_id, p.kp_id, p.sensitivity_band, distinct])
        _add_table_with_headers(doc, headers, rows)

    flips = [r for r in loo.values() if r.tier_flips or r.max_abs_shift > 0.0]
    if flips:
        doc.add_heading(
            render_text("Leave-one-out shifts", "逐一剔除影响", lang), level=2
        )
        headers = [
            render_text("KP ID", "知识点 ID", lang),
            render_text("Max |Δposterior|", "最大|Δ后验|", lang),
            render_text("Flipped year", "翻档年份", lang),
        ]
        rows = []
        for r in flips:
            flipped = ", ".join(str(y) for y in r.tier_flips) or "—"
            rows.append([r.kp_id, f"{r.max_abs_shift:.3f}", flipped])
        _add_table_with_headers(doc, headers, rows)

    warned = [p for p in posteriors if p.warnings]
    if warned:
        doc.add_heading(
            render_text("KPs with warnings", "知识点警告", lang), level=2
        )
        headers = [
            render_text("KP ID", "知识点 ID", lang),
            render_text("Warnings", "警告", lang),
        ]
        rows = [[p.kp_id, "; ".join(p.warnings)] for p in warned]
        _add_table_with_headers(doc, headers, rows)


def _render_full_kp_audit(
    doc, posteriors: list[KPPosterior], lang: ReportLang
) -> None:
    doc.add_heading(
        render_text(APPENDIX_AUDIT, "附录 A — 完整知识点审计", lang), level=1
    )
    doc.add_paragraph(
        render_text(
            "All numeric outputs for every KP. This table is for audit, not "
            "casual reading.",
            "所有知识点的完整数值。本表用于审计，非阅读用途。",
            lang,
        )
    )
    headers = [
        "kp_id",
        "tier",
        "posterior_mean",
        "ci_lower",
        "ci_upper",
        "hotness_mean",
        "hotness_std",
        "raw_hits",
        "weighted_hits",
        "lambda_used",
        "tau_used",
        "trend_label",
        "trend_delta",
        "sensitivity_band",
    ]
    rows: list[list[str]] = []
    for p in posteriors:
        rows.append(
            [
                p.kp_id,
                p.tier,
                f"{p.posterior_mean:.3f}",
                f"{p.ci_lower_95:.3f}",
                f"{p.ci_upper_95:.3f}",
                f"{p.hotness_mean_share:.3f}",
                f"{p.hotness_std_share:.3f}",
                str(p.raw_hits),
                f"{p.weighted_hits:.3f}",
                str(p.lambda_used),
                str(p.tau_used),
                p.trend_label,
                f"{p.trend_delta:.3f}",
                p.sensitivity_band,
            ]
        )
    _add_table_with_headers(doc, headers, rows)


def _render_pattern_catalogue(
    doc,
    pattern_coverage: list[PatternCoverage] | None,
    pattern_definitions: list[dict] | None,
    lang: ReportLang,
) -> None:
    doc.add_heading(
        render_text(APPENDIX_PATTERNS, "附录 B — 考法目录", lang), level=1
    )
    if not pattern_coverage and not pattern_definitions:
        doc.add_paragraph(
            render_text(
                "No pattern layer supplied for this run.",
                "本次未提供考法层。",
                lang,
            )
        )
        return
    headers = [
        "kp_id",
        "pattern_id",
        "label",
        "raw_hits",
        "weighted_hits",
        "last_seen_year",
        "saturation_index",
        "freshness_flag",
        "predicted_score",
        "tier",
    ]
    label_lookup = label_index(pattern_definitions)
    rows: list[list[str]] = []
    sorted_rows = sorted(
        pattern_coverage or [],
        key=lambda c: (c.kp_id, c.pattern_id),
    )
    for c in sorted_rows:
        rows.append(
            [
                c.kp_id,
                c.pattern_id,
                label_lookup.get(c.pattern_id, ""),
                str(c.raw_hits),
                f"{c.weighted_hits:.3f}",
                "—" if c.last_seen_year is None else f"{c.last_seen_year}",
                f"{c.saturation_index:.3f}",
                "yes" if c.freshness_flag else "no",
                f"{c.predicted_score:.3f}",
                c.tier or "—",
            ]
        )
    _add_table_with_headers(doc, headers, rows)


def _render_methodology_appendix(
    doc, hyperparameters: dict, lang: ReportLang
) -> None:
    doc.add_heading(
        render_text(APPENDIX_METHODOLOGY, "附录 C — 方法说明", lang), level=1
    )
    doc.add_paragraph(
        render_text(
            "KP layer. The per-KP probability of appearing on the next sitting "
            "is a moment-matched Beta posterior over recency-weighted hits. "
            "Each year is weighted by exp(-lambda * (reference_year - year)). "
            "The prior is Beta(tau * coverage, tau * (1 - coverage)) with tau "
            "capped at 2.0; this is a regularization prior, not an empirical "
            "one. The 95% credible interval is reported alongside the posterior "
            "mean. A (lambda, tau) sweep yields a sensitivity_band; warnings "
            "flag effective_N < 2, single-paper evidence, and all-positive / "
            "all-negative observations.",
            "知识点层：每个知识点下次出现的概率，是基于时间衰减加权命中的"
            "矩匹配 Beta 后验。每年权重为 exp(-lambda * (reference_year - year))；"
            "先验为 Beta(tau * coverage, tau * (1 - coverage))，tau 上限 2.0，"
            "属于正则化先验而非经验先验。报告 95% 可信区间。(lambda, tau) "
            "扫描产生 sensitivity_band；警告涵盖 effective_N < 2、单卷证据、"
            "全正/全负观察。",
            lang,
        )
    )
    doc.add_paragraph(
        render_text(
            "Pattern layer. Per-pattern statistics use frequency, saturation "
            "index, and a freshness flag together with a predicted score that "
            "softly rewards fresh patterns and downweights saturated ones via "
            "novelty bias alpha. No credible interval is reported at the "
            "pattern level — per-cell evidence is too sparse (typically 0–5 "
            "hits across 11–28 papers) to support an honest CI. Wording at "
            "this layer is deliberately frequency + saturation + freshness, "
            "never 'posterior'.",
            "考法层：使用频率、饱和度指数与新鲜度标志，配合 predicted_score "
            "（通过 alpha 偏置奖励未考、压低饱和考法）。考法层不给出"
            "可信区间——单格证据稀疏（常见 0–5 次/11–28 卷），"
            "任何 CI 都难以诚实表达。本层措辞坚持 频率 + 饱和度 + 新鲜度，"
            "绝不使用'后验'。",
            lang,
        )
    )
    doc.add_heading(render_text("Hyperparameters", "超参数", lang), level=2)
    headers = [render_text("Field", "字段", lang), render_text("Value", "值", lang)]
    rows = [[str(k), scalar_str(v)] for k, v in sorted(hyperparameters.items())]
    _add_table_with_headers(doc, headers, rows)
    doc.add_heading(render_text("References", "参考文献", lang), level=2)
    doc.add_paragraph("references/methodology.md", style="List Bullet")
    doc.add_paragraph("references/tier-definitions.md", style="List Bullet")


# ---------------------------------------------------------------------------
# Low-level utilities
# ---------------------------------------------------------------------------


def _set_default_font(doc) -> None:  # pragma: no cover - cosmetic
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_page_break(doc) -> None:
    """Insert a page break paragraph."""
    from docx.enum.text import WD_BREAK

    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_table_with_headers(
    doc, headers: list[str], rows: list[list[str]]
) -> None:
    """Helper that adds a styled table with one header row + data rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:  # pragma: no cover - default style fallback
        table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            if i < len(cells):
                cells[i].text = str(value) if value is not None else ""


def _normalize_sweeps(
    sweeps: dict[str, SensitivitySweep] | list,
) -> dict[str, SensitivitySweep]:
    """Accept either a kp_id-keyed dict or a flat list."""
    if isinstance(sweeps, dict):
        return sweeps
    out: dict[str, SensitivitySweep] = {}
    for sweep in sweeps:
        out[sweep.kp_id] = sweep
    return out


__all__ = ("write_docx",)
